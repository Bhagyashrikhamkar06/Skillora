import requests
import time
from docx import Document

# Create a dummy DOCX file
doc = Document()
doc.add_heading('Resume', 0)
doc.add_paragraph('Name: Test User')
doc.add_paragraph('Email: test@example.com')
doc.add_paragraph('Skills: Python, JavaScript, SQL')
doc.save('test_resume.docx')

BASE_URL = 'http://localhost:5000/api'
email = f'testuser_docx_{int(time.time())}@example.com'
password = 'TestPassword123!'

try:
    # 1. Register with correct fields
    print(f"Registering {email}...")
    reg_resp = requests.post(f'{BASE_URL}/auth/register', json={
        'full_name': 'Test User',
        'email': email,
        'password': password,
        'role': 'job_seeker'
    })
    
    if reg_resp.status_code not in [201, 409]: # 409 if already exists
        print(f"Registration failed: {reg_resp.status_code} - {reg_resp.text}")
        if reg_resp.status_code != 409:
            exit(1)

    # 2. Login
    print("Logging in...")
    login_resp = requests.post(f'{BASE_URL}/auth/login', json={'email': email, 'password': password})
    if login_resp.status_code != 200:
        print(f"Login failed: {login_resp.status_code} - {login_resp.text}")
        exit(1)
    
    token = login_resp.json()['access_token']
    print("Login successful.")

    # 3. Upload
    print("Uploading DOCX resume...")
    with open('test_resume.docx', 'rb') as f:
        files = {'file': ('test_resume.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        headers = {'Authorization': f'Bearer {token}'}
        
        resp = requests.post(f'{BASE_URL}/resume/upload', files=files, headers=headers)
        
        print(f"Status Code: {resp.status_code}")
        print(f"Response: {resp.text}")

except Exception as e:
    print(f"An error occurred: {e}")
